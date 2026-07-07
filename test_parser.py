"""
Unit tests for the biodata parser in run_match.py.

Run with:
    python -m unittest test_parser.py -v
    # or, if pytest is installed:
    pytest test_parser.py -v
"""

import os
import tempfile
import unittest

import run_match as rm


# ---------------------------------------------------------------------------
# Date of Birth parsing
# ---------------------------------------------------------------------------
class TestParseDob(unittest.TestCase):
    def test_day_monthname_year(self):
        self.assertEqual(rm.parse_dob("24 January 1998"),
                         {"day": 24, "month": 1, "year": 1998})

    def test_ordinal_suffix(self):
        self.assertEqual(rm.parse_dob("Birthdate : 3rd May 1997"),
                         {"day": 3, "month": 5, "year": 1997})

    def test_abbreviated_month(self):
        self.assertEqual(rm.parse_dob("Born 5 Sep 2000"),
                         {"day": 5, "month": 9, "year": 2000})

    def test_numeric_slash(self):
        self.assertEqual(rm.parse_dob("DOB: 24/01/1998"),
                         {"day": 24, "month": 1, "year": 1998})

    def test_numeric_dash(self):
        self.assertEqual(rm.parse_dob("07-12-1995"),
                         {"day": 7, "month": 12, "year": 1995})

    def test_missing_returns_none(self):
        self.assertIsNone(rm.parse_dob("No date here at all"))


# ---------------------------------------------------------------------------
# Time of Birth parsing
# ---------------------------------------------------------------------------
class TestParseTime(unittest.TestCase):
    def test_24h(self):
        self.assertEqual(rm.parse_time("14:16"), {"hour": 14, "minute": 16})

    def test_pm(self):
        self.assertEqual(rm.parse_time("09:43 pm"), {"hour": 21, "minute": 43})

    def test_am(self):
        self.assertEqual(rm.parse_time("2:16 AM"), {"hour": 2, "minute": 16})

    def test_noon_pm(self):
        self.assertEqual(rm.parse_time("12:48 PM"), {"hour": 12, "minute": 48})

    def test_midnight_am(self):
        self.assertEqual(rm.parse_time("12:05 AM"), {"hour": 0, "minute": 5})

    def test_labeled(self):
        self.assertEqual(rm.parse_time("Time Of Birth : 12:48 PM"),
                         {"hour": 12, "minute": 48})

    def test_missing_returns_none(self):
        self.assertIsNone(rm.parse_time("no time present"))


# ---------------------------------------------------------------------------
# Place of Birth parsing
# ---------------------------------------------------------------------------
class TestParsePlace(unittest.TestCase):
    def test_simple_label(self):
        self.assertEqual(rm.parse_place("Place of Birth: Ahmedabad"), "Ahmedabad")

    def test_birthplace_label(self):
        self.assertEqual(rm.parse_place("Birthplace: Mumbai, Maharashtra"),
                         "Mumbai, Maharashtra")

    def test_parenthetical_note_stripped(self):
        text = "Place Of Birth : Rajkot (Resided in Ahmedabad since birth)"
        self.assertEqual(rm.parse_place(text), "Rajkot")

    def test_unclosed_parenthesis_wrapped(self):
        # Simulates PDF line-wrapping inside a parenthetical
        text = "Place Of Birth : Rajkot (Resided in"
        self.assertEqual(rm.parse_place(text), "Rajkot")

    def test_missing_returns_none(self):
        self.assertIsNone(rm.parse_place("nothing relevant here"))


class TestCleanPlace(unittest.TestCase):
    def test_removes_parenthetical(self):
        self.assertEqual(rm._clean_place("Rajkot (note)"), "Rajkot")

    def test_trims_punctuation(self):
        self.assertEqual(rm._clean_place("  Ahmedabad, "), "Ahmedabad")


# ---------------------------------------------------------------------------
# Name parsing
# ---------------------------------------------------------------------------
class TestParseName(unittest.TestCase):
    def test_labeled_name(self):
        text = "PERSONAL DETAILS\nName : Vidhi Parekh\nDate Of Birth : 1 Jan 2000"
        self.assertEqual(rm.parse_name_from_text(text), "Vidhi Parekh")

    def test_does_not_pick_fathers_name(self):
        text = "Father's Name : Mr. Prashant Parekh\nName : Vidhi Parekh"
        self.assertEqual(rm.parse_name_from_text(text), "Vidhi Parekh")

    def test_skips_invocation_line(self):
        text = "|| श्री गणेशाय नम: ||\nNishant Jani"
        self.assertEqual(rm.parse_name_from_text(text), "Nishant Jani")

    def test_first_meaningful_line_fallback(self):
        text = "Bio Data\nNishant Jani\nBirthdate: 24 January 1998"
        self.assertEqual(rm.parse_name_from_text(text), "Nishant Jani")

    def test_all_caps_name(self):
        text = "SALONI SHAH\nBirthdate : 3rd May 1997"
        self.assertEqual(rm.parse_name_from_text(text), "SALONI SHAH")


class TestFirstName(unittest.TestCase):
    def test_basic(self):
        self.assertEqual(rm._first_name("Nishant Jani"), "Nishant")

    def test_all_caps_titlecased(self):
        self.assertEqual(rm._first_name("SALONI SHAH"), "Saloni")

    def test_strips_punctuation(self):
        self.assertEqual(rm._first_name("Mr. Prashant"), "Mr")

    def test_empty(self):
        self.assertEqual(rm._first_name(""), "Unknown")


# ---------------------------------------------------------------------------
# Multi-format text extraction
# ---------------------------------------------------------------------------
class TestExtractTextFromFile(unittest.TestCase):
    def test_txt(self):
        with tempfile.NamedTemporaryFile("w", suffix=".txt", delete=False,
                                         encoding="utf-8") as f:
            f.write("Name : Test User\nDOB: 24/01/1998")
            path = f.name
        try:
            text = rm.extract_text_from_file(path)
            self.assertIn("Test User", text)
        finally:
            os.unlink(path)

    def test_docx(self):
        import docx
        doc = docx.Document()
        doc.add_paragraph("Name : Docx Person")
        doc.add_paragraph("Date Of Birth : 15 August 1990")
        doc.add_paragraph("Time Of Birth : 06:30 AM")
        doc.add_paragraph("Place Of Birth : Pune")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        doc.save(path)
        try:
            text = rm.extract_text_from_file(path)
            self.assertIn("Docx Person", text)
            self.assertIn("Pune", text)
        finally:
            os.unlink(path)

    def test_unsupported_extension(self):
        with tempfile.NamedTemporaryFile("w", suffix=".rtf", delete=False) as f:
            path = f.name
        try:
            with self.assertRaises(rm.BiodataParseError) as ctx:
                rm.extract_text_from_file(path)
            self.assertIn("Unsupported file type", str(ctx.exception))
        finally:
            os.unlink(path)


# ---------------------------------------------------------------------------
# Full extract_person_info — success and clear-error behaviour
# ---------------------------------------------------------------------------
class TestExtractPersonInfo(unittest.TestCase):
    def _write_txt(self, content):
        f = tempfile.NamedTemporaryFile("w", suffix=".txt", delete=False,
                                        encoding="utf-8")
        f.write(content)
        f.close()
        self.addCleanup(os.unlink, f.name)
        return f.name

    def test_success_txt(self):
        path = self._write_txt(
            "Name : Sample Person\n"
            "Date Of Birth : 24 January 1998\n"
            "Time Of Birth : 14:16\n"
            "Place Of Birth : Ahmedabad, Gujarat\n"
        )
        info = rm.extract_person_info(path)
        self.assertEqual(info["name"], "Sample Person")
        self.assertEqual(info["day"], 24)
        self.assertEqual(info["month"], 1)
        self.assertEqual(info["year"], 1998)
        self.assertEqual(info["hour"], 14)
        self.assertEqual(info["minute"], 16)
        self.assertEqual(info["place"], "Ahmedabad, Gujarat")

    def test_missing_dob_clear_error(self):
        path = self._write_txt(
            "Name : No Date\nTime Of Birth : 14:16\nPlace Of Birth : Surat\n"
        )
        with self.assertRaises(rm.BiodataParseError) as ctx:
            rm.extract_person_info(path)
        msg = str(ctx.exception)
        self.assertIn("Date of Birth", msg)
        self.assertNotIn("Time of Birth", msg)
        self.assertNotIn("Place of Birth", msg)

    def test_missing_time_clear_error(self):
        path = self._write_txt(
            "Name : No Time\nDate Of Birth : 1 Jan 2000\nPlace Of Birth : Surat\n"
        )
        with self.assertRaises(rm.BiodataParseError) as ctx:
            rm.extract_person_info(path)
        self.assertIn("Time of Birth", str(ctx.exception))

    def test_missing_place_clear_error(self):
        path = self._write_txt(
            "Name : No Place\nDate Of Birth : 1 Jan 2000\nTime Of Birth : 10:00\n"
        )
        with self.assertRaises(rm.BiodataParseError) as ctx:
            rm.extract_person_info(path)
        self.assertIn("Place of Birth", str(ctx.exception))

    def test_multiple_missing_listed_together(self):
        path = self._write_txt("Name : Nearly Empty\n")
        with self.assertRaises(rm.BiodataParseError) as ctx:
            rm.extract_person_info(path)
        msg = str(ctx.exception)
        self.assertIn("Date of Birth", msg)
        self.assertIn("Time of Birth", msg)
        self.assertIn("Place of Birth", msg)

    def test_empty_file_clear_error(self):
        path = self._write_txt("   \n  \n")
        with self.assertRaises(rm.BiodataParseError) as ctx:
            rm.extract_person_info(path)
        self.assertIn("No readable text", str(ctx.exception))


# ---------------------------------------------------------------------------
# Optional: real sample biodata files, if present in the repo
# ---------------------------------------------------------------------------
class TestRealSamples(unittest.TestCase):
    SAMPLES = {
        "Nishant_biodata.pdf": {
            "name": "Nishant Jani", "day": 24, "month": 1, "year": 1998,
            "hour": 14, "minute": 16, "place": "Ahmedabad, Gujarat",
        },
        "SALONI_SHAH.pdf.pdf": {
            "name": "SALONI SHAH", "day": 3, "month": 5, "year": 1997,
            "hour": 21, "minute": 43, "place": "Ahmedabad, Gujarat",
        },
        "VidhiBiodataNew.pdf": {
            "name": "Vidhi Parekh", "day": 29, "month": 4, "year": 1998,
            "hour": 12, "minute": 48, "place": "Rajkot",
        },
    }

    def test_samples(self):
        for fname, expected in self.SAMPLES.items():
            if not os.path.isfile(fname):
                self.skipTest(f"sample file not present: {fname}")
            with self.subTest(file=fname):
                info = rm.extract_person_info(fname)
                for key, value in expected.items():
                    self.assertEqual(info[key], value,
                                     f"{fname}: {key} mismatch")


if __name__ == "__main__":
    unittest.main(verbosity=2)
