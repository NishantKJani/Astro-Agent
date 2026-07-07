"""
Astro Match Script
Usage: python run_match.py --boy <path_to_boy_pdf> --girl <path_to_girl_pdf>

Extracts Date of Birth, Time of Birth, and Place of Birth from each PDF,
fills the AstroSage matchmaking form, submits it, extracts the Guna score,
saves the result page as a PDF, and prints the score to stdout.
"""

import argparse
import re
import sys
import os
import asyncio
from pathlib import Path
from datetime import datetime

import pdfplumber
from playwright.async_api import async_playwright

# ---------------------------------------------------------------------------
# PDF Parsing
# ---------------------------------------------------------------------------

MONTH_MAP = {
    "january": 1, "february": 2, "march": 3, "april": 4,
    "may": 5, "june": 6, "july": 7, "august": 8,
    "september": 9, "october": 10, "november": 11, "december": 12,
    "jan": 1, "feb": 2, "mar": 3, "apr": 4,
    "jun": 6, "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}

SUPPORTED_EXTENSIONS = (".pdf", ".txt", ".doc", ".docx")


class BiodataParseError(Exception):
    """Raised when a biodata file cannot be read or required fields are missing."""


def extract_text_from_file(path: str) -> str:
    """
    Extract plain text from a biodata file.

    Supports: .pdf, .txt, .doc, .docx
    Raises BiodataParseError with a clear message for unsupported types or
    read failures.
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return _extract_text_from_pdf(path)
    if ext == ".txt":
        return _extract_text_from_txt(path)
    if ext == ".docx":
        return _extract_text_from_docx(path)
    if ext == ".doc":
        return _extract_text_from_doc(path)

    raise BiodataParseError(
        f"Unsupported file type '{ext or '(none)'}' for '{os.path.basename(path)}'.\n"
        f"Supported formats are: {', '.join(SUPPORTED_EXTENSIONS)}."
    )


def _extract_text_from_pdf(pdf_path: str) -> str:
    """Extract all text content from a PDF file."""
    try:
        text = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n".join(text)
    except Exception as e:
        raise BiodataParseError(
            f"Failed to read PDF '{os.path.basename(pdf_path)}': {e}"
        )


def _extract_text_from_txt(txt_path: str) -> str:
    """Read text from a plain-text file (tolerant of encoding issues)."""
    try:
        with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        raise BiodataParseError(
            f"Failed to read text file '{os.path.basename(txt_path)}': {e}"
        )


def _extract_text_from_docx(docx_path: str) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        import docx  # python-docx
    except ImportError:
        raise BiodataParseError(
            "Reading .docx files requires the 'python-docx' package.\n"
            "Install it with: pip install python-docx"
        )
    try:
        document = docx.Document(docx_path)
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception as e:
        raise BiodataParseError(
            f"Failed to read Word file '{os.path.basename(docx_path)}': {e}"
        )


def _extract_text_from_doc(doc_path: str) -> str:
    """
    Extract text from a legacy .doc file.

    Tries python-docx first (in case it is actually a .docx with a .doc name),
    then falls back to a best-effort extraction of readable text from the
    binary. If nothing usable is found, raises a clear error.
    """
    # Some ".doc" files are really zipped .docx — try that first.
    try:
        import docx
        document = docx.Document(doc_path)
        text = "\n".join(p.text for p in document.paragraphs)
        if text.strip():
            return text
    except Exception:
        pass

    # Best-effort: pull printable ASCII runs out of the binary stream.
    try:
        with open(doc_path, "rb") as f:
            raw = f.read()
        runs = re.findall(rb'[\x20-\x7e]{4,}', raw)
        text = "\n".join(r.decode("ascii", errors="ignore") for r in runs)
        if text.strip():
            return text
    except Exception as e:
        raise BiodataParseError(
            f"Failed to read legacy .doc file '{os.path.basename(doc_path)}': {e}"
        )

    raise BiodataParseError(
        f"Could not extract text from legacy .doc file "
        f"'{os.path.basename(doc_path)}'.\n"
        f"Please convert it to .docx, .pdf or .txt and try again."
    )


# Backwards-compatible alias
def extract_text_from_pdf(pdf_path: str) -> str:
    """Deprecated: use extract_text_from_file. Kept for compatibility."""
    return _extract_text_from_pdf(pdf_path)


def _first_name(full_name: str) -> str:
    """
    Return a filesystem-safe first name (title-cased) from a full name.

    Strips any leftover "Name:" label/bullet noise first, then scans tokens
    left-to-right for the first one that actually contains alphanumeric
    characters -- so a stray leading bullet/punctuation token can't cause
    this to silently collapse to "Unknown".
    """
    if not full_name or not full_name.strip():
        return "Unknown"

    cleaned = re.sub(
        r'(?i)^[\s' + re.escape(_BULLET_CHARS) + r']*name\s*[:\-]\s*',
        '',
        full_name.strip(),
    )

    for token in cleaned.split():
        safe = re.sub(r'[^A-Za-z0-9]', '', token)
        if safe:
            return safe.title()

    return "Unknown"


def parse_name_from_text(text: str) -> str:
    # 1. Explicit "Name:" label, tolerant of a leading bullet/marker
    label = _NAME_LABEL_RE.search(text)
    if label:
        candidate = label.group(1).strip()
        if candidate:
            return candidate

    # 2. Fallback: first meaningful line
    skip_starts = ("bio", "about", "the", "a ", "an ", "personal", "family",
                   "contact", "details")
    for raw_line in text.strip().splitlines():
        line = raw_line.strip().lstrip(_BULLET_CHARS + " \t")
        if not line:
            continue
        lower = line.lower()
        if "||" in line or line.startswith(("||", "::")):
            continue
        if lower.startswith(skip_starts):
            continue
        if line.isupper() and len(line.split()) <= 3 and ":" not in line:
            if any(w in lower for w in ("detail", "profile", "data", "about")):
                continue
        return line
    return "Unknown"


def parse_dob(text: str) -> dict:
    """
    Parse date of birth from text.
    Handles formats like:
      - "2 January 1998"
      - "3rd May 1997"
      - "Birthdate : 3rd May 1997"
      - "2/01/1998"
    Returns dict with keys: day (int), month (int), year (int),
    or None if no date could be found.
    """
    # Try "day MonthName year" pattern (with optional ordinal suffix)
    pattern1 = re.search(
        r'(\d{1,2})(?:st|nd|rd|th)?\s+([A-Za-z]+)(?:,\s*|\s+)(\d{4})',
        text, re.IGNORECASE
    )
    if pattern1:
        day = int(pattern1.group(1))
        month_str = pattern1.group(2).lower()
        year = int(pattern1.group(3))
        if month_str in MONTH_MAP:
            return {"day": day, "month": MONTH_MAP[month_str], "year": year}

    # Try numeric DD/MM/YYYY or DD-MM-YYYY
    pattern2 = re.search(r'(\d{1,2})[/\-](\d{1,2})[/\-](\d{4})', text)
    if pattern2:
        return {
            "day": int(pattern2.group(1)),
            "month": int(pattern2.group(2)),
            "year": int(pattern2.group(3)),
        }

    return None


def parse_time(text: str) -> dict:
    """
    Parse time of birth from text.
    Handles:
      - "14:16"
      - "09:43 pm" / "09:43 PM"
      - "2:16 PM"
    Returns dict with keys: hour (int 0-23), minute (int),
    or None if no time could be found.
    """
    pattern = re.search(
        r'(\d{1,2}):(\d{2})\s*(am|pm)?',
        text, re.IGNORECASE
    )
    if pattern:
        hour = int(pattern.group(1))
        minute = int(pattern.group(2))
        ampm = pattern.group(3)
        if ampm:
            ampm = ampm.lower()
            if ampm == "pm" and hour != 12:
                hour += 12
            elif ampm == "am" and hour == 12:
                hour = 0
        return {"hour": hour, "minute": minute}

    return None


def _clean_place(place: str) -> str:
    """
    Normalise a raw place string into a usable city/location.

    Removes parenthetical notes (e.g. "Rajkot (Resided in Ahmedabad...)"),
    handles wrapped/unclosed parentheses, and trims surrounding punctuation.
    """
    # Drop closed parentheticals first, then any trailing unclosed "(..."
    place = re.sub(r'\(.*?\)', '', place)
    place = re.sub(r'\(.*$', '', place)
    # Collapse whitespace and trim stray punctuation
    place = re.sub(r'\s+', ' ', place).strip(' ,.-')
    return place


def parse_place(text: str) -> str:
    """
    Extract place of birth from text.
    Handles labeled fields like "Place Of Birth :", "Birthplace:", "Place:",
    and strips parenthetical notes such as "(Resided in ... since birth)".
    """
    # Preferred: an explicit labeled field at the start of a line
    match = re.search(
        r'(?im)^\s*(?:birth\s*place|place\s*of\s*birth|birthplace|place)\s*[:\-]\s*(.+)$',
        text,
    )
    if match:
        place = _clean_place(match.group(1))
        if place:
            return place

    # Fallback: scan any line that mentions a place label
    for line in text.splitlines():
        lower = line.lower()
        if "birthplace" in lower or "place of birth" in lower or "place:" in lower or "birth city" in lower:
            if ":" in line:
                place = _clean_place(line.split(":", 1)[1])
                if place:
                    return place

    return None

_BULLET_CHARS = "•◦▪‣●○*-–—>·"

_NAME_LABEL_RE = re.compile(
    rf'(?im)^[\s{re.escape(_BULLET_CHARS)}]*Name\s*[:\-]\s*(.+)$'
)

def extract_person_info(file_path: str) -> dict:
    """
    Extract name, dob, time, and place from a biodata file (.pdf/.docx/.doc/.txt).

    Raises BiodataParseError with a clear, actionable message if the file cannot
    be read or if any required field (Date of Birth, Time of Birth, Place of
    Birth) is missing.
    """
    text = extract_text_from_file(file_path)
    fname = os.path.basename(file_path)

    if not text or not text.strip():
        raise BiodataParseError(
            f"No readable text found in '{fname}'.\n"
            f"The file may be empty, image-only/scanned, or corrupted. "
            f"Try a text-based PDF or a .txt/.docx file."
        )

    name = parse_name_from_text(text)
    dob = parse_dob(text)
    tob = parse_time(text)
    place = parse_place(text)

    # Collect every missing required field so the user sees all problems at once
    missing = []
    if dob is None:
        missing.append("Date of Birth (e.g. '24 January 1998' or '24/01/1998')")
    if tob is None:
        missing.append("Time of Birth (e.g. '14:16' or '09:43 PM')")
    if not place:
        missing.append("Place of Birth (e.g. 'Place of Birth: Ahmedabad')")

    if missing:
        details = "\n".join(f"  - {m}" for m in missing)
        raise BiodataParseError(
            f"Could not extract the following required field(s) from '{fname}':\n"
            f"{details}\n"
            f"Please make sure the file clearly contains these details."
        )

    return {
        "name": name,
        "day": dob["day"],
        "month": dob["month"],
        "year": dob["year"],
        "hour": tob["hour"],
        "minute": tob["minute"],
        "place": place,
    }


# ---------------------------------------------------------------------------
# Browser automation
# ---------------------------------------------------------------------------

async def run_matchmaking(boy: dict, girl: dict, output_pdf: str) -> str:
    """
    Open AstroSage matchmaking form in a headless browser, fill in details,
    submit, extract the Guna score, and save the result as a PDF.
    Returns the score string.
    """
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        context = await browser.new_context(
            viewport={"width": 1280, "height": 900},
            accept_downloads=True,
            user_agent=(
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            ),
        )
        page = await context.new_page()

        url = "https://www.astrosage.com/freechart/matchmaking.asp"
        print(f"[*] Navigating to {url} ...")
        await page.goto(url, wait_until="domcontentloaded", timeout=60000)
        await page.wait_for_timeout(2000)

        # --- Fill in the form ---
        print("[*] Filling Boy's details ...")
        await _fill_person_fields(page, "boy", boy)

        print("[*] Filling Girl's details ...")
        await _fill_person_fields(page, "girl", girl)

        # Step 1: Submit the main form -> goes to confirmMatchMaking.asp
        print("[*] Submitting form ...")
        try:
            async with page.expect_navigation(
                wait_until="domcontentloaded", timeout=30000
            ):
                await _submit_form(page)
        except Exception:
            await page.wait_for_timeout(3000)

        # Step 2: On the confirmation page, click "Continue" (B1) -> result page
        if "confirm" in page.url.lower():
            print("[*] Confirming details ...")
            try:
                async with page.expect_navigation(
                    wait_until="domcontentloaded", timeout=30000
                ):
                    await page.click(
                        "input[name='B1'], input[value='Continue' i]"
                    )
            except Exception:
                await page.wait_for_timeout(3000)

        # Wait for results page to load
        print("[*] Waiting for results ...")
        try:
            await page.wait_for_selector(
                "text=/total.*point|guna.*milan|out of 36|received.*out of/i",
                timeout=30000,
            )
        except Exception:
            pass  # page may have loaded anyway

        await page.wait_for_timeout(2000)

        # Extract score
        score = await _extract_score(page)

        # Download the official, nicely-formatted Match Making PDF report
        await _download_match_pdf(page, context, output_pdf)

        await browser.close()
        return score


async def _fill_person_fields(page, person_type: str, info: dict):
    """
    Fill in the form fields for a person (boy or girl).
    AstroSage matchmaking form uses suffix '1' for the boy and '2' for the girl:
      name1/name2, day1/day2, month1/month2, year1/year2,
      hrs1/hrs2, min1/min2, sec1/sec2, place1/place2
    """
    n = "1" if person_type == "boy" else "2"

    # Name (text input)
    await page.fill(f"input[name='name{n}']", info["name"])

    # Day dropdown (option value = day number, no zero padding)
    await _select_dropdown(page, f"select[name='day{n}']", str(info["day"]))

    # Month dropdown (option value = month number 1-12)
    await _select_dropdown(page, f"select[name='month{n}']", str(info["month"]))

    # Year (text input)
    await page.fill(f"input[name='year{n}']", str(info["year"]))

    # Hour dropdown (option value = 0-23)
    await _select_dropdown(page, f"select[name='hrs{n}']", str(info["hour"]))

    # Minute dropdown (option value = 0-59, no zero padding)
    await _select_dropdown(page, f"select[name='min{n}']", str(info["minute"]))

    # Seconds dropdown -> 0
    await _select_dropdown(page, f"select[name='sec{n}']", "0")

    # Place of birth (autocomplete input)
    await _fill_place_with_autocomplete(
        page, f"input[name='place{n}']", info["place"], person_type
    )


async def _select_dropdown(page, selector: str, value: str):
    """Select a value in a <select> dropdown, trying multiple strategies."""
    try:
        # Try selecting by value
        await page.select_option(selector, value=value)
        return
    except Exception:
        pass
    try:
        # Try selecting by label (for month names)
        await page.select_option(selector, label=value)
        return
    except Exception:
        pass
    # Zero-padded fallback
    try:
        await page.select_option(selector, value=value.zfill(2))
    except Exception:
        pass


async def _fill_place_with_autocomplete(page, selector: str, place: str, person_type: str):
    """
    Fill a place field that shows a jQuery-UI autocomplete dropdown.
    Types the city name and clicks the first matching suggestion so the site
    sets the hidden city id / coordinates needed for the calculation.
    """
    # Use just the primary city token (before the first comma) for matching
    city = place.split(",")[0].strip()
    try:
        await page.click(selector)
        await page.fill(selector, "")
        await page.type(selector, city, delay=120)

        # Wait for the jQuery UI autocomplete list to appear
        try:
            await page.wait_for_selector(
                "ul.ui-autocomplete li.ui-menu-item:visible", timeout=8000
            )
        except Exception:
            pass
        await page.wait_for_timeout(500)

        suggestions = await page.query_selector_all(
            "ul.ui-autocomplete li.ui-menu-item"
        )
        # Filter to visible suggestions
        for sug in suggestions:
            try:
                if await sug.is_visible():
                    await sug.click()
                    await page.wait_for_timeout(500)
                    return
            except Exception:
                continue

        # Fallback: press ArrowDown + Enter to select first suggestion
        await page.focus(selector)
        await page.keyboard.press("ArrowDown")
        await page.keyboard.press("Enter")
        await page.wait_for_timeout(500)
    except Exception as e:
        print(f"  [!] Warning: could not fill place for {person_type}: {e}")


async def _submit_form(page):
    """Click the submit button on the matchmaking form."""
    submit_selectors = [
        "input[name='submit']",
        "input[type='submit']",
        "button[type='submit']",
        "input[value='Submit' i]",
        "input[value='Match' i]",
        "button:has-text('Submit')",
        "button:has-text('Match')",
    ]
    for sel in submit_selectors:
        try:
            btn = await page.query_selector(sel)
            if btn:
                await btn.click()
                return
        except Exception:
            continue
    raise RuntimeError("Could not find submit button on matchmaking form.")


async def _extract_score(page) -> str:
    """
    Extract the Guna score from the result page.
    AstroSage typically shows something like "Total Points: 28/36"
    """
    content = await page.content()

    # Pattern: look for a number out of 36
    patterns = [
        r'(\d+(?:\.\d+)?)\s*/\s*36',               # "28/36" or "28.5/36"
        r'total\s+points?\s*[:\-]?\s*(\d+(?:\.\d+)?)',  # "Total Points: 28"
        r'(\d+(?:\.\d+)?)\s+out\s+of\s+36',         # "28 out of 36"
        r'guna\s+milan\s*[:\-]?\s*(\d+(?:\.\d+)?)',  # "Guna Milan: 28"
        r'score\s*[:\-]?\s*(\d+(?:\.\d+)?)',          # "Score: 28"
    ]
    for pat in patterns:
        m = re.search(pat, content, re.IGNORECASE)
        if m:
            score_val = m.group(1)
            return f"{score_val}/36"

    # Try to get text from common result elements
    result_selectors = [
        ".match-score",
        "#matchScore",
        ".guna-score",
        "td:has-text('36')",
        "span:has-text('/36')",
    ]
    for sel in result_selectors:
        try:
            el = await page.query_selector(sel)
            if el:
                text = await el.inner_text()
                m = re.search(r'(\d+(?:\.\d+)?)', text)
                if m:
                    return f"{m.group(1)}/36"
        except Exception:
            continue

    return "Score not found (please check the saved PDF)"


async def _download_match_pdf(page, context, output_pdf: str):
    """
    Click the official "Download Match Making PDF" button on the result page and
    save the downloaded report to output_pdf. Falls back to fetching the PDF
    href directly, and finally to a page screenshot PDF if needed.
    """
    link_selector = "a:has-text('Download Match Making PDF')"

    # Strategy 1: click the link and capture the browser download
    try:
        link = await page.query_selector(link_selector)
        if link:
            try:
                async with page.expect_download(timeout=45000) as dl_info:
                    await link.click()
                download = await dl_info.value
                await download.save_as(output_pdf)
                print(f"[*] Match Making PDF downloaded to: {output_pdf}")
                return
            except Exception:
                # The link may open the PDF in a new tab instead of downloading
                href = await link.get_attribute("href")
                if href:
                    await _fetch_pdf_via_request(context, href, output_pdf)
                    return
    except Exception as e:
        print(f"  [!] Could not click download button: {e}")

    # Strategy 2: build/fetch the PDF URL directly from the href
    try:
        link = await page.query_selector(link_selector)
        href = await link.get_attribute("href") if link else None
        if href:
            await _fetch_pdf_via_request(context, href, output_pdf)
            return
    except Exception as e:
        print(f"  [!] Could not fetch PDF via href: {e}")

    # Strategy 3: last-resort fallback — render the page itself to PDF
    print("  [!] Falling back to page-render PDF.")
    await page.pdf(path=output_pdf, format="A4", print_background=True)
    print(f"[*] Result saved to: {output_pdf}")


async def _fetch_pdf_via_request(context, href: str, output_pdf: str):
    """Download a PDF directly using the browser context's request API."""
    resp = await context.request.get(href, timeout=60000)
    body = await resp.body()
    with open(output_pdf, "wb") as f:
        f.write(body)
    print(f"[*] Match Making PDF downloaded to: {output_pdf}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Astro Match: Compare two kundlis and get the Guna score."
    )
    parser.add_argument(
        "--boy", required=True,
        help="Path to boy's biodata file (.pdf, .docx, .doc or .txt)",
    )
    parser.add_argument(
        "--girl", required=True,
        help="Path to girl's biodata file (.pdf, .docx, .doc or .txt)",
    )
    parser.add_argument(
        "--output",
        default=None,
        help=(
            "Output PDF file name. Defaults to "
            "<BoyFirstName>_<GirlFirstName>_astro_match_result.pdf"
        ),
    )
    args = parser.parse_args()

    boy_file = os.path.abspath(args.boy)
    girl_file = os.path.abspath(args.girl)

    if not os.path.isfile(boy_file):
        print(f"[ERROR] Boy's biodata file not found: {boy_file}", file=sys.stderr)
        sys.exit(1)
    if not os.path.isfile(girl_file):
        print(f"[ERROR] Girl's biodata file not found: {girl_file}", file=sys.stderr)
        sys.exit(1)

    # Parse biodata files with clear error reporting
    try:
        print("[*] Parsing boy's biodata ...")
        boy = extract_person_info(boy_file)
        print(f"    Name : {boy['name']}")
        print(f"    DOB  : {boy['day']:02d}/{boy['month']:02d}/{boy['year']}")
        print(f"    Time : {boy['hour']:02d}:{boy['minute']:02d}")
        print(f"    Place: {boy['place']}")
    except BiodataParseError as e:
        print(f"\n[ERROR] Problem with the BOY's biodata file:\n{e}\n", file=sys.stderr)
        sys.exit(2)

    try:
        print("[*] Parsing girl's biodata ...")
        girl = extract_person_info(girl_file)
        print(f"    Name : {girl['name']}")
        print(f"    DOB  : {girl['day']:02d}/{girl['month']:02d}/{girl['year']}")
        print(f"    Time : {girl['hour']:02d}:{girl['minute']:02d}")
        print(f"    Place: {girl['place']}")
    except BiodataParseError as e:
        print(f"\n[ERROR] Problem with the GIRL's biodata file:\n{e}\n", file=sys.stderr)
        sys.exit(2)


    # Determine output file name (default uses both first names)
    if args.output:
        output_pdf = os.path.abspath(args.output)
    else:
        boy_first = _first_name(boy["name"])
        girl_first = _first_name(girl["name"])
        output_pdf = os.path.abspath(
            f"{boy_first}_{girl_first}_astro_match_result.pdf"
        )

    # Run browser automation
    score = asyncio.run(run_matchmaking(boy, girl, output_pdf))

    print("\n" + "=" * 50)
    print(f"  ASTRO MATCH SCORE : {score}")
    print("=" * 50 + "\n")


if __name__ == "__main__":
    main()